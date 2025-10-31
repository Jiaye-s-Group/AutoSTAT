import ast
import io
import re
from io import BytesIO

import streamlit as st
from tqdm import tqdm
from stqdm import stqdm
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px
import plotly.io as pio
from concurrent.futures import ThreadPoolExecutor, as_completed

from utils.sanitize_code import sanitize_code
from workflow.report.report_core import *


def report_prepare(agents, parallel=True, max_workers=4):
    report_agent = agents[-1]
    toc = report_agent.load_outline()
    if toc is None:
        st.error("Please generate the table of contents first.")
        return

    toc = sanitize_code(toc)

    # === Summarize abstracts from each analysis module ===
    agent_abstracts = {}
    with st.spinner("Summarizing results from each analysis module..."):
        for i in stqdm(range(len(agents) - 1)):
            agent_abstracts[i] = agents[i].check_abstract()

    # === Update TOC with FIG list ===
    selected_full_contents_vis = agents[2].check_full()
    toc = report_agent.selected_photo_update_toc(toc, selected_full_contents_vis)
    toc = sanitize_code(toc)
    print(toc)
    try:
        toc = ast.literal_eval(toc)
    except Exception:
        pass

    # === Update TOC with module selection list ===
    with st.spinner("Matching required analysis modules for each section..."):
        toc_with_choice = report_agent.update_toc_with_relevant_sections(toc, agent_abstracts)
        toc_with_choice = sanitize_code(toc_with_choice)
        try:
            toc_with_choice = ast.literal_eval(toc_with_choice)
        except Exception:
            pass

    # === Initialize report structure ===
    doc = Reportcore()
    doc.add_heading('Data Analysis Report', 0)
    selected_model = st.session_state.selected_model

    def process_section(idx, t, t_w_c, history_content=""):
        st.session_state.selected_model = selected_model
        # t: ('Title', Level, Content Outline, [figs], [modules])
        _, _, _, _, choice_list = t_w_c
        selected_full_contents = {i: agents[i].check_full() for i in choice_list if i < len(agents) - 1}
        content = report_agent.write_section_body(toc, t, selected_full_contents, history_content)
        print(idx)
        return (idx, t, content)

    results = []

    # Serial or parallel processing
    if not parallel:
        with st.spinner("Generating chapter content serially (with context)..."):
            history_content = ""
            for idx, t in stqdm(enumerate(toc)):
                t_w_c = toc_with_choice[idx]
                _, _, content = process_section(idx, t, t_w_c, history_content)
                results.append((idx, t, content))
                history_content += f"\n\n{t[0]}\n{content}"
    else:
        with st.spinner(f"Generating chapter content in parallel ({max_workers} threads)..."):
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                print(toc_with_choice)
                # print(f"idx={idx}, len={len(toc_with_choice)}")
                futures = {
                    executor.submit(process_section, idx, t, toc_with_choice[idx], ""): idx
                    for idx, t in enumerate(toc)
                }
                for future in stqdm(as_completed(futures), total=len(futures)):
                    try:
                        results.append(future.result())
                    except Exception as e:
                        print(f"Chapter generation failed: {e}")

    # Sort & write to report
    results.sort(key=lambda x: x[0])
    for _, t, content in results:
        doc.add_heading(t[0], level=t[1])
        doc.add_paragraph(content)

    report_agent.save_report(doc)
