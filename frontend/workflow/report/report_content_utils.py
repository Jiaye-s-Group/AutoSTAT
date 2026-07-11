import base64
import binascii
import html
import io
import json
import re
import unicodedata
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

WORD_EXPORT_IMAGE_SCALE = 0.6
WORD_EXPORT_IMAGE_BASE_WIDTH_INCHES = 6.0
WORD_EXPORT_IMAGE_WIDTH_INCHES = WORD_EXPORT_IMAGE_BASE_WIDTH_INCHES * WORD_EXPORT_IMAGE_SCALE
WORD_EXPORT_FONT_LATIN = "Times New Roman"
WORD_EXPORT_FONT_EAST_ASIA = "Microsoft YaHei"
WORD_EXPORT_FONT_MONO = "Consolas"
WORD_EXPORT_TEXT_COLOR = RGBColor(31, 41, 55)
WORD_EXPORT_HEADING_COLOR = RGBColor(17, 24, 39)
WORD_EXPORT_MUTED_COLOR = RGBColor(107, 114, 128)
WORD_EXPORT_CODE_BACKGROUND = "F5F7FA"
WORD_EXPORT_QUOTE_BACKGROUND = "F8FAFC"
FIG_PLACEHOLDER_CORE_PATTERN = r"[\[\uFF3B\u3010]?\s*FIG\s*[:\uFF1A]?\s*\d+\s*[\]\uFF3D\u3011]?"
FIG_PLACEHOLDER_CAPTURE_PATTERN = r"(?<![A-Za-z0-9_])[\[\uFF3B\u3010]?\s*FIG\s*[:\uFF1A]?\s*(\d+)\s*[\]\uFF3D\u3011]?(?![A-Za-z0-9_])"
FIG_PLACEHOLDER_PATTERN = rf"(?<![A-Za-z0-9_]){FIG_PLACEHOLDER_CORE_PATTERN}(?![A-Za-z0-9_])"
FIG_TRAILING_PUNCTUATION_PATTERN = re.compile(
    rf"(?<![A-Za-z0-9_])(?P<placeholders>{FIG_PLACEHOLDER_CORE_PATTERN}(?:\s*{FIG_PLACEHOLDER_CORE_PATTERN})*)(?![A-Za-z0-9_])(?P<spacing>\s*)(?P<punctuation>[\u3002\uFF01\uFF1F\uFF1B\uFF0C\u3001\uFF1A]+)",
    flags=re.IGNORECASE,
)
FIG_NUMBER_TOKEN_PATTERN = r"[0-9０-９零〇一二两三四五六七八九十百千①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]+"
FIG_LABEL_PATTERN = r"(?:[FfＦｆ][IiＩｉ][GgＧｇ](?:\.|ure)?|[Ff][Ii][Gg](?:\.|ure)?|图表|图片|插图|图)"
FIG_REFERENCE_CAPTURE_RE = re.compile(
    rf"(?<![A-Za-z0-9_])[\[\uFF3B\u3010(（]?\s*"
    rf"{FIG_LABEL_PATTERN}\s*(?:[:：#＃.\-—_ ]\s*)?"
    rf"(?P<num>{FIG_NUMBER_TOKEN_PATTERN})\s*[\]\uFF3D\u3011)）]?"
    rf"(?![A-Za-z0-9_])",
    flags=re.IGNORECASE,
)
FIG_REVERSE_REFERENCE_CAPTURE_RE = re.compile(
    rf"(?<![A-Za-z0-9_])[\[\uFF3B\u3010(（]?\s*第\s*"
    rf"(?P<num>{FIG_NUMBER_TOKEN_PATTERN})\s*(?:张|幅|个|份)?\s*"
    rf"(?:图表|图片|插图|图)\s*[\]\uFF3D\u3011)）]?"
    rf"(?![A-Za-z0-9_])",
    flags=re.IGNORECASE,
)
FIG_PLACEHOLDER_CAPTURE_RE = re.compile(FIG_PLACEHOLDER_CAPTURE_PATTERN, flags=re.IGNORECASE)
CHINESE_NUMERAL_VALUES = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "二": 2,
    "两": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
}
INLINE_HEADING_BODY_STARTERS = (
    "本章",
    "本节",
    "本部分",
    "本报告",
    "本文",
    "首先",
    "其次",
    "然后",
    "接下来",
    "随后",
    "最后",
    "此外",
    "同时",
    "通过",
    "针对",
    "为了",
    "基于",
    "这里",
    "因此",
    "其中",
    "需要",
)
INLINE_HEADING_BODY_PATTERN = re.compile(
    rf"^(?P<title>.+?[。！？!?])\s*(?P<body>(?:{'|'.join(map(re.escape, INLINE_HEADING_BODY_STARTERS))}).+)$"
)
INLINE_MARKDOWN_HEADING_AFTER_TEXT_PATTERN = re.compile(
    r"^(?P<prefix>.*?[。！？!?；;：:])\s*(?P<heading>#{1,6}[ \t\u3000]*.+)$"
)


TABLE_TITLE_PATTERN = re.compile(r"^(?:表|Table)\s*\d+[\s\u3000:：].+$", flags=re.IGNORECASE)


def maybe_json_loads(value: Any) -> Any:
    if not isinstance(value, str):
        return value

    stripped = value.strip()
    if not stripped:
        return value

    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        return value


def find_nested_field(data: Any, field_name: str) -> Any:
    if isinstance(data, dict):
        if field_name in data:
            return data[field_name]

        for nested_value in data.values():
            nested = find_nested_field(nested_value, field_name)
            if nested is not None:
                return nested

    if isinstance(data, list):
        for item in data:
            nested = find_nested_field(item, field_name)
            if nested is not None:
                return nested

    return None


def find_first_nested_field(data: Any, field_names: list[str]) -> Any:
    for field_name in field_names:
        value = find_nested_field(data, field_name)
        if value is not None:
            return value
    return None


def stringify_string(value: Any) -> str:
    value = maybe_json_loads(value)

    if value is None:
        return ""

    if isinstance(value, str):
        return value.strip()

    return json.dumps(value, ensure_ascii=False, indent=2)


def normalize_toc_list(value: Any) -> list[str]:
    parsed_value = maybe_json_loads(value)

    if isinstance(parsed_value, list):
        normalized_items: list[str] = []
        for item in parsed_value:
            item_text = str(item).replace("\\r\\n", "\n").replace("\\n", "\n").strip()
            if not item_text:
                continue
            normalized_items.extend([line.strip() for line in item_text.splitlines() if line.strip()])
        return normalized_items

    if isinstance(parsed_value, str):
        normalized_text = parsed_value.replace("\\r\\n", "\n").replace("\\n", "\n")
        return [line.strip() for line in normalized_text.splitlines() if line.strip()]

    return []


def extract_report_markdown(result: Any) -> str:
    if isinstance(result, str):
        return result.strip()

    candidate = find_first_nested_field(
        result,
        [
            "report_markdown",
            "markdown",
            "report_md",
            "md",
            "report_content",
            "content",
            "report",
            "body",
        ],
    )
    return stringify_string(candidate)


def extract_report_html(result: Any) -> str:
    candidate = find_first_nested_field(result, ["final_html", "report_html", "html"])
    if isinstance(candidate, list):
        return "".join(str(item) for item in candidate if item is not None).strip()
    return stringify_string(candidate)


def extract_report_text(result: Any) -> str:
    candidate = find_first_nested_field(
        result,
        ["report_text", "text", "report_content", "content", "report", "body"],
    )
    return stringify_string(candidate)


def extract_report_word_bytes(result: Any) -> bytes | None:
    candidate = find_first_nested_field(
        result,
        ["report_word", "word", "report_word_base64", "word_base64", "docx_base64"],
    )

    if isinstance(candidate, (bytes, bytearray)):
        return bytes(candidate)

    if not isinstance(candidate, str):
        return None

    stripped = candidate.strip()
    if not stripped:
        return None

    try:
        return base64.b64decode(stripped, validate=True)
    except (binascii.Error, ValueError):
        return None


def _parse_chinese_figure_number(text: str) -> int | None:
    if not text:
        return None

    total = 0
    section = 0
    current = 0
    seen = False
    unit_map = {"十": 10, "百": 100, "千": 1000}

    for char in text:
        if char in CHINESE_NUMERAL_VALUES:
            current = CHINESE_NUMERAL_VALUES[char]
            seen = True
            continue

        unit = unit_map.get(char)
        if unit is None:
            return None

        if current == 0:
            current = 1
        section += current * unit
        current = 0
        seen = True

    if not seen:
        return None

    total += section + current
    return total


def _parse_figure_number_token(value: Any) -> int | None:
    token = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not token:
        return None

    if token.isdigit():
        return int(token)

    try:
        numeric = unicodedata.numeric(token)
        if float(numeric).is_integer():
            return int(numeric)
    except (TypeError, ValueError):
        pass

    return _parse_chinese_figure_number(token)


def normalize_figure_placeholders(text: str) -> str:
    """Normalize common figure references to `[FIG:n]` placeholders."""
    if not isinstance(text, str) or not text:
        return text

    def replace(match: re.Match[str]) -> str:
        number = _parse_figure_number_token(match.group("num"))
        if number is None:
            return match.group(0)
        return f"[FIG:{number}]"

    normalized = FIG_REVERSE_REFERENCE_CAPTURE_RE.sub(replace, text)
    normalized = FIG_REFERENCE_CAPTURE_RE.sub(replace, normalized)
    return normalized


def remove_figure_placeholders(text: str) -> str:
    """Remove standardized and Chinese-compatible figure placeholders from prose."""
    if not isinstance(text, str) or not text:
        return text

    normalized = normalize_figure_placeholders(text)
    cue_words = r"(?:(?:另|并)?(?:如|见|参见|详见|参考|根据|结合|从))?"
    suffix_words = r"(?:所示|可见|可以看出|显示|展示)?"
    phrase_pattern = re.compile(
        rf"{cue_words}\s*{FIG_PLACEHOLDER_CAPTURE_PATTERN}\s*{suffix_words}\s*[，,、:：；;]?",
        flags=re.IGNORECASE,
    )
    normalized = phrase_pattern.sub("", normalized)
    normalized = FIG_PLACEHOLDER_CAPTURE_RE.sub("", normalized)
    normalized = re.sub(r"\s+([，,。.!！?？；;：:、])", r"\1", normalized)
    normalized = re.sub(r"[，,、；;：:]\s*([。.!！?？])", r"\1", normalized)
    normalized = re.sub(r"([（(【\[])\s+([）)】\]])", "", normalized)
    normalized = re.sub(r"[ \t]{2,}", " ", normalized)
    normalized = re.sub(r"\n[ \t]+", "\n", normalized)
    return normalized.strip()


def normalize_trailing_punctuation_before_figure_placeholder(text: str) -> str:
    if not isinstance(text, str) or not text.strip():
        return text

    return FIG_TRAILING_PUNCTUATION_PATTERN.sub(
        lambda match: f"{match.group('punctuation')}{match.group('spacing')}{match.group('placeholders')}",
        text,
    )


def _split_inline_heading_content(text: str) -> tuple[str, str | None]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return "", None

    body_starters = (*INLINE_HEADING_BODY_STARTERS, "\u672c\u9636\u6bb5", "\u8be5\u9636\u6bb5")

    spaced_split = re.match(r"^(?P<title>\S(?:.*?\S)?)\s{2,}(?P<body>\S.*)$", text.strip())
    if spaced_split:
        return spaced_split.group("title").strip(), spaced_split.group("body").strip()

    starter_split = re.match(
        rf"^(?P<title>.+?)\s+(?P<body>(?:{'|'.join(map(re.escape, body_starters))}).+)$",
        normalized,
    )
    if starter_split:
        return starter_split.group("title").strip(), starter_split.group("body").strip()

    inline_split = INLINE_HEADING_BODY_PATTERN.match(normalized)
    if inline_split:
        return inline_split.group("title").strip(), inline_split.group("body").strip()

    return normalized, None


def _parse_markdown_heading_line_with_level(text: str) -> tuple[int, str, str | None] | None:
    if not isinstance(text, str):
        return None

    normalized = text.replace("\\r\\n", "\n").replace("\\n", "\n").strip()
    if not normalized:
        return None

    heading_match = re.match(r"^(#{1,6})[ \t\u3000]*(.*)$", normalized)
    if not heading_match:
        return None

    heading_level = len(heading_match.group(1))
    heading_text = heading_match.group(2).strip()
    if not heading_text:
        return None

    title_text, body_text = _split_inline_heading_content(heading_text)
    if not title_text:
        return None
    return heading_level, title_text, body_text


def _parse_markdown_heading_line(text: str) -> tuple[str, str | None] | None:
    parsed = _parse_markdown_heading_line_with_level(text)
    if parsed is None:
        return None
    _, title_text, body_text = parsed
    return title_text, body_text


def _docx_heading_level_from_html(html_level: int) -> int:
    return max(1, min(6, html_level - 1 if html_level > 1 else 1))


def _docx_heading_level_from_markdown(markdown_level: int) -> int:
    return _docx_heading_level_from_html(markdown_level)


def _split_text_with_markdown_headings(text: str) -> list[tuple[str, str]]:
    normalized = text.strip()
    if not normalized:
        return []

    parsed_heading = _parse_markdown_heading_line(normalized)
    if parsed_heading:
        heading_text, body_text = parsed_heading
        segments: list[tuple[str, str]] = [("heading", heading_text)]
        if body_text:
            segments.append(("text", body_text))
        return segments

    inline_heading_match = INLINE_MARKDOWN_HEADING_AFTER_TEXT_PATTERN.match(normalized)
    if inline_heading_match:
        prefix_text = inline_heading_match.group("prefix").strip()
        heading_segments = _split_text_with_markdown_headings(inline_heading_match.group("heading"))
        if heading_segments:
            segments = []
            if prefix_text:
                segments.append(("text", prefix_text))
            segments.extend(heading_segments)
            return segments

    return [("text", normalized)]


def _split_markdown_heading_lines(text: str) -> list[tuple[str, str]]:
    if not isinstance(text, str):
        return []

    normalized = text.replace("\\r\\n", "\n").replace("\\n", "\n")
    lines = [line.strip() for line in normalized.split("\n") if line.strip()]
    if not lines:
        return []

    parsed_lines: list[tuple[str, str]] = []
    has_heading = False
    for line in lines:
        segments = _split_text_with_markdown_headings(line)
        parsed_lines.extend(segments)
        has_heading = has_heading or any(line_kind == "heading" for line_kind, _ in segments)

    return parsed_lines if has_heading else []


def _parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return []

    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]

    cells = [part.strip().replace("\\|", "|") for part in re.split(r"(?<!\\)\|", stripped)]
    if not any(cells):
        return []
    return cells


def _is_markdown_table_delimiter_row(cells: list[str], expected_columns: int | None = None) -> bool:
    if not cells:
        return False
    if expected_columns is not None and len(cells) != expected_columns:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _extract_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int] | None:
    if start_index + 1 >= len(lines):
        return None

    header_cells = _parse_markdown_table_row(lines[start_index])
    delimiter_cells = _parse_markdown_table_row(lines[start_index + 1])
    if not header_cells or not _is_markdown_table_delimiter_row(delimiter_cells, expected_columns=len(header_cells)):
        return None

    rows: list[list[str]] = [header_cells]
    next_index = start_index + 2

    while next_index < len(lines):
        stripped = lines[next_index].strip()
        if not stripped:
            break

        row_cells = _parse_markdown_table_row(stripped)
        if not row_cells or _is_markdown_table_delimiter_row(row_cells):
            break
        rows.append(row_cells)
        next_index += 1

    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    return normalized_rows, next_index


def _markdown_table_rows_to_html(rows: list[list[str]]) -> str:
    if not rows:
        return ""

    header_cells = "".join(f"<th>{html.escape(cell)}</th>" for cell in rows[0])
    body_rows = rows[1:]
    body_html = "".join(
        "<tr>" + "".join(f"<td>{html.escape(cell)}</td>" for cell in row) + "</tr>"
        for row in body_rows
    )
    table_html = (
        '<table class="report-model-comparison-table">'
        "<thead><tr>"
        f"{header_cells}"
        "</tr></thead>"
    )
    if body_html:
        table_html += f"<tbody>{body_html}</tbody>"
    table_html += "</table>"
    return table_html


def _looks_like_table_title(text: str) -> bool:
    normalized = str(text or "").strip()
    return bool(normalized and TABLE_TITLE_PATTERN.match(normalized))


def html_to_markdown(html_text: str) -> str:
    if not html_text.strip():
        return ""

    html_text = normalize_trailing_punctuation_before_figure_placeholder(html_text)

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup
    lines: list[str] = []

    block_tags = {
        "article",
        "aside",
        "div",
        "main",
        "section",
        "p",
        "ul",
        "ol",
        "li",
        "figure",
    }

    def table_to_markdown(table_tag: Tag) -> str:
        def escape_cell(text: str) -> str:
            return str(text or "").replace("|", "\\|")

        row_tags = table_tag.find_all("tr")
        rows: list[list[str]] = []
        for row_tag in row_tags:
            cell_tags = row_tag.find_all(["th", "td"], recursive=False)
            if not cell_tags:
                cell_tags = row_tag.find_all(["th", "td"])
            row_values = [cell.get_text(" ", strip=True) for cell in cell_tags]
            if row_values:
                rows.append(row_values)

        if not rows:
            return ""

        column_count = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
        header = normalized_rows[0]
        separator = ["---"] * column_count
        body = normalized_rows[1:] if len(normalized_rows) > 1 else []
        lines = [
            "| " + " | ".join(escape_cell(cell) for cell in header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        lines.extend("| " + " | ".join(escape_cell(cell) for cell in row) + " |" for row in body)
        return "\n".join(lines).strip()

    def add_line(text: str) -> None:
        clean = text.strip()
        if clean:
            lines.append(clean)

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                add_line(text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                add_line(f"{'#' * int(node.name[1])} {text}")
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                add_line(f"- {text}")
            return

        if node.name == "img":
            img_src = (node.get("src") or "").strip()
            alt_text = (node.get("alt") or "图表").strip()
            if img_src:
                add_line(f"![{alt_text}]({img_src})")
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            img = node.find("img")
            img_src = (img.get("src") or "").strip() if img else ""
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else ""
            if img_src:
                add_line(f"![{caption_text or '图表'}]({img_src})")
            return

        if node.name == "table":
            table_markdown = table_to_markdown(node)
            if table_markdown:
                add_line(table_markdown)
            return

        if node.name == "p":
            has_media_children = any(
                isinstance(child, Tag)
                and (
                    child.name == "img"
                    or child.name == "figure"
                    or (child.name == "div" and "report-figure-block" in (child.get("class") or []))
                )
                for child in node.children
            )
            if has_media_children:
                for child in node.children:
                    walk(child)
                return
            text = node.get_text(" ", strip=True)
            if text:
                add_line(text)
            return

        if node.name in block_tags:
            child_tags = [child for child in node.children if isinstance(child, Tag)]
            has_block_children = any(
                (
                    child.name in block_tags
                    or (child.name and child.name.startswith("h") and len(child.name) == 2 and child.name[1].isdigit())
                    or child.name == "img"
                )
                for child in child_tags
            )
            if not has_block_children:
                text = node.get_text(" ", strip=True)
                if text:
                    add_line(text)
                return

            direct_text = " ".join(
                text.strip()
                for text in node.find_all(string=True, recursive=False)
                if text and text.strip()
            )
            if direct_text:
                add_line(direct_text)
            for child in node.children:
                walk(child)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    return "\n\n".join(lines).strip()


def build_markdown_preview_from_html(html_text: str, max_chars: int = 20000) -> str:
    if not html_text.strip():
        return ""

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup
    blocks: list[str] = []

    def append_block(text: str) -> None:
        clean = re.sub(r"\s+", " ", text).strip()
        if clean:
            blocks.append(clean)

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                append_block(text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name == "br":
            blocks.append("")
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                append_block(f"{'#' * int(node.name[1])} {text}")
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                append_block(f"- {text}")
            return

        if node.name == "img":
            append_block("![图表](embedded-image)")
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else ""
            append_block(f"![{caption_text or '图表'}](embedded-image)")
            return

        if node.name in {"p", "figcaption"}:
            text = node.get_text(" ", strip=True)
            if text:
                append_block(text)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    preview = "\n\n".join(blocks)
    preview = re.sub(
        r"(?:\s*!\[图表\]\(embedded-image\)\s*){2,}",
        "\n\n![图表](embedded-image)\n\n",
        preview,
        flags=re.IGNORECASE,
    )
    preview = re.sub(r"\n{3,}", "\n\n", preview).strip()

    if len(preview) > max_chars:
        preview = preview[:max_chars].rstrip() + "\n\n...[预览已截断，下载文件保留完整内容]"

    return preview


def _configure_doc_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = WORD_EXPORT_FONT_LATIN
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(8)

    title_style = doc.styles["Title"]
    title_style.font.name = WORD_EXPORT_FONT_LATIN
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = WORD_EXPORT_HEADING_COLOR
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(18)
    title_style.paragraph_format.keep_with_next = True

    heading_sizes = {
        "Heading 1": 18,
        "Heading 2": 15,
        "Heading 3": 13,
        "Heading 4": 12,
        "Heading 5": 11,
        "Heading 6": 11,
    }
    for style_name, font_size in heading_sizes.items():
        heading_style = doc.styles[style_name]
        heading_style.font.name = WORD_EXPORT_FONT_LATIN
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
        heading_style.font.size = Pt(font_size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = WORD_EXPORT_HEADING_COLOR
        heading_style.paragraph_format.space_before = Pt(16 if style_name == "Heading 1" else 14)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        list_style = doc.styles[style_name]
        list_style.font.name = WORD_EXPORT_FONT_LATIN
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
        list_style.font.size = Pt(11)
        list_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
        list_style.paragraph_format.line_spacing = 1.35
        list_style.paragraph_format.space_after = Pt(4)

    code_style = _get_or_create_paragraph_style(doc, "ChatGPT Code")
    code_style.base_style = normal_style
    code_style.font.name = WORD_EXPORT_FONT_MONO
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_MONO)
    code_style.font.size = Pt(9.5)
    code_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.05)
    code_style.paragraph_format.line_spacing = 1.15
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    quote_style = _get_or_create_paragraph_style(doc, "ChatGPT Quote")
    quote_style.base_style = normal_style
    quote_style.font.name = WORD_EXPORT_FONT_LATIN
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    quote_style.font.size = Pt(10.5)
    quote_style.font.italic = True
    quote_style.font.color.rgb = WORD_EXPORT_MUTED_COLOR
    quote_style.paragraph_format.left_indent = Inches(0.25)
    quote_style.paragraph_format.right_indent = Inches(0.05)
    quote_style.paragraph_format.line_spacing = 1.35
    quote_style.paragraph_format.space_before = Pt(6)
    quote_style.paragraph_format.space_after = Pt(6)

    caption_style = _get_or_create_paragraph_style(doc, "ChatGPT Caption")
    caption_style.base_style = normal_style
    caption_style.font.name = WORD_EXPORT_FONT_LATIN
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    caption_style.font.size = Pt(9.5)
    caption_style.font.italic = True
    caption_style.font.color.rgb = WORD_EXPORT_MUTED_COLOR
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.line_spacing = 1.2
    caption_style.paragraph_format.space_before = Pt(0)
    caption_style.paragraph_format.space_after = Pt(10)

    table_title_style = _get_or_create_paragraph_style(doc, "ChatGPT Table Title")
    table_title_style.base_style = normal_style
    table_title_style.font.name = WORD_EXPORT_FONT_LATIN
    table_title_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    table_title_style.font.size = Pt(10.5)
    table_title_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
    table_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_title_style.paragraph_format.line_spacing = 1.2
    table_title_style.paragraph_format.space_before = Pt(6)
    table_title_style.paragraph_format.space_after = Pt(6)


def _get_or_create_paragraph_style(doc: Document, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def _set_run_font(
    run,
    font_name: str = WORD_EXPORT_FONT_LATIN,
    east_asia_font: str = WORD_EXPORT_FONT_EAST_ASIA,
    size: Pt | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _style_paragraph_runs(
    paragraph,
    font_name: str = WORD_EXPORT_FONT_LATIN,
    east_asia_font: str = WORD_EXPORT_FONT_EAST_ASIA,
) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, east_asia_font=east_asia_font)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _set_paragraph_bottom_border(paragraph, color: str = "5B7DB1", size: str = "12", space: str = "1") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)

    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_table_borders(
    table,
    *,
    top: dict[str, str] | None = None,
    bottom: dict[str, str] | None = None,
    left: dict[str, str] | None = None,
    right: dict[str, str] | None = None,
    inside_h: dict[str, str] | None = None,
    inside_v: dict[str, str] | None = None,
) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)
    tbl_borders = OxmlElement("w:tblBorders")

    border_map = {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
        "insideH": inside_h,
        "insideV": inside_v,
    }
    for edge_name, border_value in border_map.items():
        edge = OxmlElement(f"w:{edge_name}")
        border_value = border_value or {"val": "nil"}
        for attr_name, attr_value in border_value.items():
            edge.set(qn(f"w:{attr_name}"), attr_value)
        tbl_borders.append(edge)

    tbl_pr.append(tbl_borders)


def _set_cell_borders(
    cell,
    *,
    top: dict[str, str] | None = None,
    bottom: dict[str, str] | None = None,
    left: dict[str, str] | None = None,
    right: dict[str, str] | None = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is not None:
        tc_pr.remove(tc_borders)
    tc_borders = OxmlElement("w:tcBorders")

    border_map = {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }
    for edge_name, border_value in border_map.items():
        edge = OxmlElement(f"w:{edge_name}")
        border_value = border_value or {"val": "nil"}
        for attr_name, attr_value in border_value.items():
            edge.set(qn(f"w:{attr_name}"), attr_value)
        tc_borders.append(edge)

    tc_pr.append(tc_borders)


def _apply_three_line_table_style(table, header_rows: set[int], total_rows: int) -> None:
    border = {"val": "single", "sz": "12", "color": "111827", "space": "0"}
    _set_table_borders(
        table,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )

    for row_index, row in enumerate(table.rows):
        is_header = row_index in header_rows
        is_last_row = row_index == total_rows - 1
        for cell in row.cells:
            _set_cell_borders(
                cell,
                top=border if is_header else {"val": "nil"},
                bottom=border if is_header or is_last_row else {"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def _add_table_title(doc: Document, text: str):
    paragraph = _add_text_paragraph(doc, text, style_name="ChatGPT Table Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def _add_text_paragraph(doc: Document, text: str, style_name: str = "Normal"):
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.add_run(text)
    _style_paragraph_runs(paragraph)
    return paragraph


def _add_heading_paragraph(
    doc: Document,
    text: str,
    level: int,
    state: dict[str, bool],
    allow_title_style: bool = False,
) -> None:
    style_name = (
        "Title"
        if allow_title_style and level == 1 and not state["has_title"]
        else f"Heading {min(level, 6)}"
    )
    paragraph = _add_text_paragraph(doc, text, style_name=style_name)
    paragraph.paragraph_format.keep_with_next = True
    if style_name == "Title":
        paragraph.paragraph_format.space_after = Pt(12)
        _set_paragraph_bottom_border(paragraph)
    state["has_title"] = True


def _add_body_paragraph(doc: Document, text: str):
    return _add_text_paragraph(doc, text, style_name="Normal")


def _add_list_item(doc: Document, text: str, ordered: bool = False):
    style_name = "List Number" if ordered else "List Bullet"
    return _add_text_paragraph(doc, text, style_name=style_name)


def _add_caption(doc: Document, text: str):
    paragraph = _add_text_paragraph(doc, text, style_name="ChatGPT Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def _add_quote_block(doc: Document, text: str) -> None:
    paragraph = _add_text_paragraph(doc, text, style_name="ChatGPT Quote")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    _set_paragraph_shading(paragraph, WORD_EXPORT_QUOTE_BACKGROUND)


def _add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="ChatGPT Code")
    lines = text.rstrip("\n").splitlines() or [text]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line if line else " ")
        _set_run_font(
            run,
            font_name=WORD_EXPORT_FONT_MONO,
            east_asia_font=WORD_EXPORT_FONT_MONO,
            size=Pt(9.5),
            color=WORD_EXPORT_TEXT_COLOR,
        )
        if index < len(lines) - 1:
            run.add_break()
    _set_paragraph_shading(paragraph, WORD_EXPORT_CODE_BACKGROUND)


def _add_table_from_rows(doc: Document, rows: list[list[str]], header_rows: set[int] | None = None) -> bool:
    if not rows:
        return False

    header_rows = header_rows or {0}
    column_count = max(len(row) for row in rows)
    if column_count == 0:
        return False

    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized_rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row_values in enumerate(normalized_rows):
        for column_index in range(column_count):
            text = row_values[column_index]
            cell = table.cell(row_index, column_index)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                _style_paragraph_runs(paragraph)
                if row_index in header_rows:
                    for run in paragraph.runs:
                        run.font.bold = True

    _apply_three_line_table_style(table, header_rows=header_rows, total_rows=len(normalized_rows))
    doc.add_paragraph("")
    return True


def _add_table_from_html(doc: Document, table_tag: Tag) -> bool:
    row_tags = table_tag.find_all("tr")
    if not row_tags:
        return False

    rows: list[list[str]] = []
    header_rows: set[int] = set()

    for row_tag in row_tags:
        cell_tags = row_tag.find_all(["th", "td"], recursive=False)
        if not cell_tags:
            cell_tags = row_tag.find_all(["th", "td"])
        if not cell_tags:
            continue
        rows.append([cell.get_text(" ", strip=True) for cell in cell_tags])
        if any(cell.name == "th" for cell in cell_tags):
            header_rows.add(len(rows) - 1)

    return _add_table_from_rows(doc, rows, header_rows=header_rows or {0})


def _add_docx_image(doc: Document, image_buffer: io.BytesIO, caption_text: str | None = None) -> None:
    image_buffer.seek(0)
    doc.add_picture(image_buffer, width=Inches(WORD_EXPORT_IMAGE_WIDTH_INCHES))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4 if caption_text else 10)
    if caption_text:
        _add_caption(doc, caption_text)


def _decode_data_image_uri(image_src: str) -> io.BytesIO | None:
    if not image_src.startswith("data:image"):
        return None

    try:
        _, encoded = image_src.split(",", 1)
        image_bytes = base64.b64decode(encoded)
    except (ValueError, binascii.Error):
        return None

    return io.BytesIO(image_bytes)


def build_docx_from_html(html_text: str) -> bytes:
    html_text = normalize_trailing_punctuation_before_figure_placeholder(html_text)

    doc = Document()
    _configure_doc_style(doc)

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup

    block_tags = {
        "article",
        "aside",
        "blockquote",
        "div",
        "figure",
        "figcaption",
        "main",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "li",
        "ul",
    }
    state = {"has_title": False}

    def add_image_from_tag(img_tag: Tag, caption_text: str | None = None) -> bool:
        image_src = (img_tag.get("src") or "").strip()
        image_buffer = _decode_data_image_uri(image_src)
        if image_buffer is None:
            return False
        _add_docx_image(doc, image_buffer, caption_text=caption_text)
        return True

    def add_markdown_text(text: str) -> bool:
        lines = [line.strip() for line in text.replace("\\r\\n", "\n").replace("\\n", "\n").split("\n") if line.strip()]
        if not lines:
            return False

        parsed_by_line = [_parse_markdown_heading_line_with_level(line) for line in lines]
        if not any(parsed is not None for parsed in parsed_by_line):
            return False

        for line in lines:
            parsed = _parse_markdown_heading_line_with_level(line)
            if parsed is None:
                _add_body_paragraph(doc, line)
                continue
            markdown_level, line_text, body_text = parsed
            _add_heading_paragraph(
                doc,
                line_text,
                level=_docx_heading_level_from_markdown(markdown_level),
                state=state,
            )
            if body_text:
                _add_body_paragraph(doc, body_text)
        return True

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                if not add_markdown_text(text):
                    _add_body_paragraph(doc, text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name == "br":
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                html_level = int(node.name[1])
                _add_heading_paragraph(
                    doc,
                    text,
                    level=_docx_heading_level_from_html(html_level),
                    state=state,
                    allow_title_style=(html_level == 1),
                )
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                _add_list_item(doc, text, ordered=(node.parent is not None and node.parent.name == "ol"))
            return

        if node.name == "img":
            add_image_from_tag(node)
            return

        if node.name == "figure":
            img_tag = node.find("img")
            caption_tag = node.find("figcaption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else None
            if img_tag is not None:
                add_image_from_tag(img_tag, caption_text=caption_text)
            elif caption_text:
                _add_caption(doc, caption_text)
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            img_tag = node.find("img")
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else None
            if img_tag is not None:
                add_image_from_tag(img_tag, caption_text=caption_text)
            elif caption_text:
                _add_caption(doc, caption_text)
            return

        if node.name == "div" and "report-modeling-table-block" in (node.get("class") or []):
            title_tag = node.find(class_="report-modeling-table-title")
            table_tag = node.find("table")
            title_text = title_tag.get_text(" ", strip=True) if title_tag else ""
            if title_text:
                _add_table_title(doc, title_text)
            if table_tag is not None:
                _add_table_from_html(doc, table_tag)
            return

        if node.name == "figcaption":
            text = node.get_text(" ", strip=True)
            if text:
                _add_caption(doc, text)
            return

        if node.name == "p" and "report-modeling-table-title" in (node.get("class") or []):
            text = node.get_text(" ", strip=True)
            if text:
                _add_table_title(doc, text)
            return

        if node.name == "blockquote":
            text = node.get_text("\n", strip=True)
            if text:
                _add_quote_block(doc, text)
            return

        if node.name == "pre":
            text = node.get_text("\n", strip=False).strip("\n")
            if text:
                _add_code_block(doc, text)
            return

        if node.name == "table":
            _add_table_from_html(doc, node)
            return

        if node.name == "p":
            has_media_children = any(
                isinstance(child, Tag)
                and (
                    child.name == "img"
                    or child.name == "figure"
                    or (child.name == "div" and "report-figure-block" in (child.get("class") or []))
                )
                for child in node.children
            )
            if has_media_children:
                for child in node.children:
                    if isinstance(child, Tag) and child.name == "img":
                        add_image_from_tag(child)
                    else:
                        walk(child)
                return
            direct_images = [child for child in node.children if isinstance(child, Tag) and child.name == "img"]
            if direct_images and not node.get_text(" ", strip=True):
                for img_tag in direct_images:
                    add_image_from_tag(img_tag)
                return
            text = node.get_text("\n", strip=True)
            if text:
                if add_markdown_text(text):
                    return
                _add_body_paragraph(doc, text)
            return

        if node.name == "hr":
            doc.add_paragraph("")
            return

        if node.name in block_tags:
            direct_text = "\n".join(
                text.strip()
                for text in node.find_all(string=True, recursive=False)
                if text and text.strip()
            )
            handled_direct_text = False
            if direct_text and node.name not in {"ol", "ul"}:
                if not add_markdown_text(direct_text):
                    _add_body_paragraph(doc, direct_text)
                handled_direct_text = True
            for child in node.children:
                if isinstance(child, Tag) and child.name == "img":
                    add_image_from_tag(child)
                elif isinstance(child, Tag):
                    walk(child)
                elif not handled_direct_text:
                    walk(child)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def markdown_to_html(markdown_text: str, title: str = "") -> str:
    """
    轻量 Markdown -> HTML：
    - 不生成目录
    - 不生成 aside
    - 只保留正文结构，交给 report_render 补充导出样式
    """
    body_parts: list[str] = []
    in_ul = False
    lines = markdown_text.splitlines()

    def close_list() -> None:
        nonlocal in_ul
        if in_ul:
            body_parts.append("</ul>")
            in_ul = False

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            close_list()
            index += 1
            continue

        if _looks_like_table_title(stripped):
            parsed_table = _extract_markdown_table(lines, index + 1)
            if parsed_table is not None:
                close_list()
                table_rows, next_index = parsed_table
                body_parts.append(
                    "<div class='report-modeling-table-block'>"
                    f"<p class='report-modeling-table-title'>{html.escape(stripped)}</p>"
                    f"{_markdown_table_rows_to_html(table_rows)}"
                    "</div>"
                )
                index = next_index
                continue

        parsed_table = _extract_markdown_table(lines, index)
        if parsed_table is not None:
            close_list()
            table_rows, next_index = parsed_table
            body_parts.append(_markdown_table_rows_to_html(table_rows))
            index = next_index
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            close_list()
            level = len(heading_match.group(1))
            text = html.escape(heading_match.group(2).strip())
            body_parts.append(f"<h{level}>{text}</h{level}>")
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            if not in_ul:
                body_parts.append("<ul>")
                in_ul = True
            body_parts.append(f"<li>{html.escape(bullet_match.group(1).strip())}</li>")
            index += 1
            continue

        close_list()
        body_parts.append(f"<p>{html.escape(stripped)}</p>")
        index += 1

    close_list()
    body_html = "\n".join(body_parts)

    page_title = html.escape(title or "")
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{page_title}</title>
</head>
<body>
  <main class="report-body">
    {body_html}
  </main>
</body>
</html>"""


def build_docx_from_markdown(markdown_text: str) -> bytes:
    markdown_text = normalize_trailing_punctuation_before_figure_placeholder(markdown_text)

    doc = Document()
    _configure_doc_style(doc)
    state = {"has_title": False}
    code_lines: list[str] = []
    in_code_block = False
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped_line = line.strip()

        if stripped_line.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if not stripped_line:
            index += 1
            continue

        if _looks_like_table_title(stripped_line):
            parsed_table = _extract_markdown_table(lines, index + 1)
            if parsed_table is not None:
                table_rows, next_index = parsed_table
                _add_table_title(doc, stripped_line)
                _add_table_from_rows(doc, table_rows, header_rows={0})
                index = next_index
                continue

        parsed_table = _extract_markdown_table(lines, index)
        if parsed_table is not None:
            table_rows, next_index = parsed_table
            _add_table_from_rows(doc, table_rows, header_rows={0})
            index = next_index
            continue

        parsed_segments = _split_markdown_heading_lines(stripped_line)
        if parsed_segments:
            for line_kind, line_text in parsed_segments:
                if line_kind == "heading":
                    _add_heading_paragraph(doc, line_text, level=1, state=state)
                else:
                    _add_body_paragraph(doc, line_text)
            index += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped_line)
        if ordered_match:
            _add_list_item(doc, ordered_match.group(1).strip(), ordered=True)
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped_line)
        if bullet_match:
            _add_list_item(doc, bullet_match.group(1).strip(), ordered=False)
            index += 1
            continue

        quote_match = re.match(r"^>\s?(.*)$", stripped_line)
        if quote_match:
            _add_quote_block(doc, quote_match.group(1).strip())
            index += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\((.+)\)$", stripped_line)
        if image_match:
            image_buffer = _decode_data_image_uri(image_match.group(2).strip())
            if image_buffer is not None:
                caption_text = image_match.group(1).strip() or None
                _add_docx_image(doc, image_buffer, caption_text=caption_text)
            index += 1
            continue

        _add_body_paragraph(doc, stripped_line)
        index += 1

    if in_code_block and code_lines:
        _add_code_block(doc, "\n".join(code_lines))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def truncate_text(value: Any, max_chars: int = 4000) -> str:
    if value is None:
        return ""

    text = stringify_string(value)
    if not text:
        return ""

    text = text.strip()
    if len(text) <= max_chars:
        return text

    return text[:max_chars] + "\n\n...[内容过长，已截断]"


VISUAL_TOC_MAX_TOPICS = 12
VISUAL_TOC_ANALYSES_PER_TOPIC = 3
VISUAL_TOC_ITEM_ANALYSIS_CHARS = 260
VISUAL_TOC_TOPIC_ANALYSIS_CHARS = 700


def _single_line_text(value: Any, max_chars: int) -> str:
    text = stringify_string(value)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + "...[截断]"


def _plotly_title_text(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        text = value.get("text") or value.get("title") or value.get("name")
        if isinstance(text, dict):
            return _plotly_title_text(text)
        if text is not None:
            return str(text).strip()
    return ""


def _extract_plotly_metadata(fig_value: Any) -> dict[str, Any]:
    fig = maybe_json_loads(fig_value)
    if not isinstance(fig, dict):
        return {}

    layout = fig.get("layout") if isinstance(fig.get("layout"), dict) else {}
    data = fig.get("data") if isinstance(fig.get("data"), list) else []

    trace_names: list[str] = []
    trace_types: list[str] = []
    for trace in data:
        if not isinstance(trace, dict):
            continue
        trace_name = str(trace.get("name", "")).strip()
        trace_type = str(trace.get("type", "")).strip()
        if trace_name and trace_name not in trace_names:
            trace_names.append(trace_name)
        if trace_type and trace_type not in trace_types:
            trace_types.append(trace_type)

    return {
        "title": _plotly_title_text(layout.get("title")),
        "x_axis": _plotly_title_text((layout.get("xaxis") or {}).get("title") if isinstance(layout.get("xaxis"), dict) else ""),
        "y_axis": _plotly_title_text((layout.get("yaxis") or {}).get("title") if isinstance(layout.get("yaxis"), dict) else ""),
        "legend": _plotly_title_text((layout.get("legend") or {}).get("title") if isinstance(layout.get("legend"), dict) else ""),
        "trace_names": trace_names[:8],
        "trace_types": trace_types[:4],
    }


def _clean_visual_topic_token(value: Any, max_chars: int = 36) -> str:
    text = stringify_string(value)
    text = html.unescape(text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" ：:，,。.;；、-_/\\|")
    text = re.sub(
        r"^(?:图表|图中|该图|可视化结果|结果|数据显示|"
        r"小提琴图|散点图矩阵|散点图|热力图|平行坐标图|PCA降维图|"
        r"箱线图|柱状图|条形图|折线图|直方图|饼图|雷达图)"
        r"(?:展示|显示|反映|呈现|说明|揭示)?(?:了)?",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?:频率)?(?:分布)?(?:直方图|柱状图|条形图|折线图|散点图|箱线图|热力图|饼图|雷达图|图表|可视化|分析)$",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = text.strip(" 的：:，,。.;；、-_/\\|")

    generic_terms = {
        "id", "index", "idx", "count", "value", "variable",
        "图表", "数据", "变量", "字段", "数值", "样本", "类别",
        "分布", "关系", "趋势", "比较", "分析", "主要变量", "整体数据",
    }
    if not text or text.lower() in generic_terms:
        return ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip()


def _topic_from_analysis_text(analysis: str) -> str:
    relation_match = re.search(
        r"([\u4e00-\u9fffA-Za-z0-9_]{2,20}?)(?:与|和|及|、)([\u4e00-\u9fffA-Za-z0-9_]{2,20}?)(?:之间)?(?:的)?(?:关系|相关|差异|比较|对比)",
        analysis,
        flags=re.IGNORECASE,
    )
    if relation_match:
        topic = _clean_visual_topic_token(f"{relation_match.group(1)}与{relation_match.group(2)}")
        if topic:
            return topic

    trend_match = re.search(
        r"([\u4e00-\u9fffA-Za-z0-9_]{2,24}?)(?:随|在|按)([\u4e00-\u9fffA-Za-z0-9_]{1,12}?)(?:变化|呈现|呈|上升|下降|波动|趋势)",
        analysis,
        flags=re.IGNORECASE,
    )
    if trend_match:
        topic = _clean_visual_topic_token(f"{trend_match.group(1)}随{trend_match.group(2)}")
        if topic:
            return topic

    patterns = (
        r"([A-Za-z_][A-Za-z0-9_]*(?:\s*(?:与|和|及|、|,|and|vs\.?)\s*[A-Za-z_][A-Za-z0-9_]*)+)",
        r"([\u4e00-\u9fffA-Za-z0-9_]{2,30}?)(?:的)?(?:分布|趋势|变化|差异|占比|相关|关系|重要性|表现)",
    )
    for pattern in patterns:
        match = re.search(pattern, analysis, flags=re.IGNORECASE)
        if match:
            topic = _clean_visual_topic_token(match.group(1))
            if topic:
                return topic

    first_sentence = re.split(r"[。！？!?；;\n]", analysis, maxsplit=1)[0]
    return _clean_visual_topic_token(first_sentence, max_chars=28)


def _visual_topic_from_fig_item(item: Any, index: int) -> tuple[str, str]:
    analysis = ""
    metadata: dict[str, Any] = {}

    if isinstance(item, dict):
        analysis = stringify_string(item.get("analysis") or item.get("desc") or "")
        metadata = _extract_plotly_metadata(item.get("fig"))
    else:
        analysis = stringify_string(item)

    x_axis = _clean_visual_topic_token(metadata.get("x_axis"))
    y_axis = _clean_visual_topic_token(metadata.get("y_axis"))
    title = _clean_visual_topic_token(metadata.get("title"))
    legend = _clean_visual_topic_token(metadata.get("legend"))
    trace_names = []
    for name in metadata.get("trace_names", []):
        clean_name = _clean_visual_topic_token(name)
        if clean_name:
            trace_names.append(clean_name)

    if x_axis and y_axis and x_axis.lower() != y_axis.lower():
        topic = f"{x_axis}与{y_axis}"
    else:
        topic = y_axis or x_axis or title or legend
        if not topic and trace_names:
            topic = "、".join(trace_names[:3])
        if not topic:
            topic = _topic_from_analysis_text(analysis)
        if not topic:
            topic = f"图表{index + 1}"

    return topic, analysis


def _compress_fig_analysis_for_toc(fig_analysis: list[Any]) -> list[dict[str, Any]]:
    groups: dict[str, dict[str, Any]] = {}

    for index, item in enumerate(fig_analysis):
        topic, analysis = _visual_topic_from_fig_item(item, index)
        key = re.sub(r"\s+", "", topic).lower()
        if key not in groups:
            groups[key] = {
                "topic": topic,
                "figure_count": 0,
                "analyses": [],
            }

        group = groups[key]
        group["figure_count"] += 1
        if analysis and len(group["analyses"]) < VISUAL_TOC_ANALYSES_PER_TOPIC:
            group["analyses"].append(_single_line_text(analysis, VISUAL_TOC_ITEM_ANALYSIS_CHARS))

    topics = list(groups.values())
    if len(topics) > VISUAL_TOC_MAX_TOPICS:
        kept_topics = topics[: VISUAL_TOC_MAX_TOPICS - 1]
        overflow_topics = topics[VISUAL_TOC_MAX_TOPICS - 1 :]
        kept_topics.append(
            {
                "topic": "其他可视化主题",
                "figure_count": sum(int(topic.get("figure_count", 0)) for topic in overflow_topics),
                "analyses": [
                    "还覆盖以下主题："
                    + _single_line_text("、".join(str(topic.get("topic", "")) for topic in overflow_topics), 600)
                ],
            }
        )
        topics = kept_topics

    compressed: list[dict[str, Any]] = []
    for topic in topics:
        analyses = [text for text in topic.get("analyses", []) if text]
        analysis_text = "；".join(analyses)
        compressed.append(
            {
                "topic": topic.get("topic", ""),
                "figure_count": topic.get("figure_count", 0),
                "finding": _single_line_text(analysis_text, VISUAL_TOC_TOPIC_ANALYSIS_CHARS),
            }
        )

    return compressed


def shrink_summary_for_toc(summary: Any) -> dict[str, Any]:
    """
    给 Reporting_toc 用的 summary 瘦身版：
    - 保留 title / desc / result / df / processed_df 的短文本
    - fig_analysis 会先按变量/主题聚合，再保留压缩后的主题摘要
    - 丢掉大 code / 大 full / 超长内容
    """
    if not isinstance(summary, dict):
        return {}

    out: dict[str, Any] = {}

    if "title" in summary:
        out["title"] = truncate_text(summary.get("title"), 200)

    if "desc" in summary:
        out["desc"] = truncate_text(summary.get("desc"), 1000)

    if "result" in summary:
        out["result"] = truncate_text(summary.get("result"), 1000)

    if "df" in summary:
        out["df"] = truncate_text(summary.get("df"), 600)

    if "processed_df" in summary:
        out["processed_df"] = truncate_text(summary.get("processed_df"), 600)

    fig_analysis = summary.get("fig_analysis")
    if isinstance(fig_analysis, list) and fig_analysis:
        out["figure_count"] = len(fig_analysis)
        out["fig_analysis"] = _compress_fig_analysis_for_toc(fig_analysis)

    return out


def normalize_toc_md_input(value: Any) -> list[str]:
    """
    toc_md 可能是 list，也可能是字符串，这里统一成 list[str]
    """
    if value is None:
        return []

    if isinstance(value, str):
        text = value.replace("\\r\\n", "\n").replace("\\n", "\n")
        return [line.strip() for line in text.split("\n") if line.strip()]

    if isinstance(value, list):
        out = []
        for item in value:
            item_text = str(item).strip()
            if item_text:
                out.append(item_text)
        return out

    return []

def normalize_part(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_for_dedup(text: str) -> str:
    normalized = normalize_part(text)
    normalized = re.sub(r"\s+", "", normalized)
    return normalized


def sanitize_section_heading_text(text: str) -> str:
    normalized = str(text or "").strip()
    if not normalized:
        return ""

    normalized = re.sub(r"[（(][^()（）]*[）)]", "", normalized)
    normalized = re.sub(r"[【\[][^][】\[]*[】\]]", "", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def wrap_section_as_markdown(section: Any, content: str) -> str:
    content = normalize_part(content)

    title = ""
    level = 2

    if isinstance(section, dict):
        num = str(section.get("num", "")).strip()
        raw_title = sanitize_section_heading_text(section.get("title", ""))
        level_raw = section.get("level", 1)

        try:
            level = int(level_raw)
        except Exception:
            level = 1

        title = f"{num} {raw_title}".strip()
    else:
        title = sanitize_section_heading_text(section)
        if re.match(r"^\d+\.\d+\.\d+\s+", title):
            level = 3
        elif re.match(r"^\d+\.\d+\s+", title):
            level = 2
        else:
            level = 1

    heading_prefix = {
        1: "##",
        2: "###",
        3: "####",
        4: "#####",
    }.get(level, "###")

    if not content:
        return f"{heading_prefix} {title}\n"
    return f"{heading_prefix} {title}\n\n{content}\n"


def build_history_context(history_parts: list[str], max_chars: int = 1800) -> str:
    """
    history 只给 writer 看，不进入最终成品拼接。
    只保留最近一段，避免 prompt 越滚越长。
    """
    if not history_parts:
        return ""

    joined = "\n\n".join(
        normalize_part(part)
        for part in history_parts
        if normalize_part(part)
    ).strip()

    if len(joined) <= max_chars:
        return joined

    return joined[-max_chars:]
